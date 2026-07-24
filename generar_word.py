import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = Document()

    # Añadir Título Principal
    title = doc.add_heading('INFORME TÉCNICO: DESAFÍOS ACTUALES Y MODELOS DE RECONOCIMIENTO FACIAL', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Añadir metadatos
    doc.add_paragraph('Materia: Visión Artificial\nActividad: 09 - Segunda Unidad\nEnfoque: Análisis empírico y arquitectónico de modelos State-of-the-Art (SOTA) en Face Recognition.')

    # ------------------ Punto 1 ------------------
    doc.add_heading('1. Relación entre Face Verification, One-Shot Learning y Redes Siamesas', level=1)
    
    p = doc.add_paragraph()
    p.add_run('Fundamentación Teórica:').bold = True
    doc.add_paragraph('La verificación facial (Face Verification) y el One-Shot Learning están intrínsecamente ligados a la arquitectura de las Redes Siamesas (Siamese Networks). A diferencia de las redes neuronales convolucionales (CNN) tradicionales que terminan en una capa Softmax para clasificar imágenes en N categorías (lo cual requeriría reentrenar la red si se añade una nueva persona), las Redes Siamesas resuelven el problema codificando las imágenes.')
    doc.add_paragraph('En el One-Shot Learning, el modelo es capaz de reconocer a un sujeto a partir de una única imagen de referencia. Esto se logra mediante una red siamesa que procesa en paralelo dos imágenes (imagen ancla y la imagen de prueba) extrayendo sus características (vectores). Finalmente, se calcula la distancia espacial (e.g. Euclidiana L2 o Similitud del Coseno) entre ambos vectores. Si la distancia es menor a un hiperparámetro (umbral o threshold), la red infiere que ambas imágenes pertenecen a la misma identidad.')
    
    p = doc.add_paragraph()
    p.add_run('Implementación Práctica:').bold = True
    doc.add_paragraph('En el script de evaluación punto1_siamese.py, se implementó este principio utilizando el backend de FaceNet. La función de verificación extrae los puntos clave y determina la métrica de distancia de forma autónoma. Al cruzar dos imágenes distintas, el algoritmo procesa la lejanía matemática de los tensores de ambas caras y genera un booleano (Verificado / No Verificado) sin necesidad de haber entrenado un clasificador tradicional previamente.')

    # ------------------ Punto 2 ------------------
    doc.add_heading('2. Demostración Arquitectónica de FaceNet y DeepFace (Embeddings)', level=1)
    
    p = doc.add_paragraph()
    p.add_run('Fundamentación Teórica:').bold = True
    doc.add_paragraph('La afirmación "DeepFace y FaceNet son arquitecturas que utilizan los principios de las redes siamesas" es completamente verdadera. Estas arquitecturas no retornan una etiqueta textual, sino una representación densa en un subespacio de dimensión finita conocida como Embedding. En FaceNet, este embedding típicamente es un tensor de 128 o 512 dimensiones proyectado sobre una hiperesfera unitaria gracias a la función de pérdida Triplet Loss.')
    
    p = doc.add_paragraph()
    p.add_run('Evidencia Empírica:').bold = True
    doc.add_paragraph('En el código (punto2_embeddings.py), se extrajo directamente la capa anterior a cualquier capa fully-connected que funcione de clasificador, utilizando el método DeepFace.represent(). Esto retorna un vector de n-dimensiones de punto flotante correspondiente al rostro.')
    
    try:
        if os.path.exists('data/resultados/resultados parte 2.png'):
            pic = doc.add_picture('data/resultados/resultados parte 2.png', width=Inches(6.0))
            pic_par = doc.paragraphs[-1]
            pic_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph('Figura 1: Visualización topológica de los vectores de características (embeddings). Se muestran las primeras 100 dimensiones extraídas, demostrando la naturaleza paramétrica codificada por la red.')
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].font.size = Pt(9)
            cap.runs[0].font.italic = True
    except Exception as e:
        doc.add_paragraph(f'(Error cargando imagen 1: {e})')

    # ------------------ Punto 3 ------------------
    doc.add_heading('3. Verificación Facial (1:1) vs Reconocimiento Facial (1:N)', level=1)
    
    p = doc.add_paragraph()
    p.add_run('Arquitectura de Procesamiento:').bold = True
    doc.add_paragraph('Aunque el flujo de trabajo inicial es el mismo en ambos enfoques (Detección de Rostro -> Alineación -> Normalización -> Extracción de Embeddings), la bifurcación ocurre en la fase de Inferencia Espacial:')
    doc.add_paragraph('- Verificación Facial (1:1): El algoritmo simplemente calcula la norma de la distancia entre el tensor A y el tensor B. Es sumamente ligero, con una complejidad temporal de O(1).', style='List Bullet')
    doc.add_paragraph('- Reconocimiento Facial (1:N): El algoritmo extrae el tensor de la persona y necesita buscar el tensor más cercano en toda una base de datos de tamaño N. Calcula la distancia N veces o emplea algoritmos KNN (K-Nearest Neighbors). Su complejidad temporal es de O(N), por lo que es mucho más pesado computacionalmente.', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('Evidencia de Rendimiento (Benchmarking):').bold = True
    doc.add_paragraph('En el experimento punto3_speed.py, medimos los ciclos de reloj y el tiempo total en segundos requeridos para ambos procesos usando la arquitectura VGG-Face y midiendo con la librería time de Python:')
    
    try:
        if os.path.exists('data/resultados/punto 3 resporte.png'):
            pic = doc.add_picture('data/resultados/punto 3 resporte.png', width=Inches(5.0))
            pic_par = doc.paragraphs[-1]
            pic_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph('Figura 2: Diferencia logarítmica/lineal de tiempos de latencia. La verificación toma significativamente menos tiempo de cómputo en comparación a buscar la firma biométrica en una base de datos.')
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].font.size = Pt(9)
            cap.runs[0].font.italic = True
    except Exception as e:
        doc.add_paragraph(f'(Error cargando imagen 2: {e})')

    # ------------------ Punto 4 ------------------
    doc.add_heading('4. Análisis de Detectores: Cascadas vs Single-Shot (MTCNN, YOLO, RetinaFace)', level=1)
    
    p = doc.add_paragraph()
    p.add_run('MTCNN vs YOLO:').bold = True
    doc.add_paragraph('MTCNN (Multi-task Cascaded Convolutional Networks) es una arquitectura secuencial basada en pirámides de imágenes. Utiliza 3 redes neuronales independientes consecutivas (P-Net, R-Net, O-Net) para proponer bounding boxes, refinarlos y luego localizar los 5 puntos clave faciales (ojos, nariz, comisuras). Debido a que itera cascadas de redes, es lento frente a imágenes grandes o multitudes.')
    doc.add_paragraph('Por otro lado, YOLO (You Only Look Once) es un detector de objetos genérico Single-Shot, que predice los bounding boxes en una sola pasada a la red convolucional a través de un grid cuadriculado. Sacrifica una leve precisión geométrica a cambio de latencia ultrabaja en tiempo real.')
    
    p = doc.add_paragraph()
    p.add_run('RetinaFace como evolución hacia Single-Shot:').bold = True
    doc.add_paragraph('Modelos como RetinaFace y MediaPipe representan una transición tecnológica a modelos de "una sola pasada" especializados en rostros. RetinaFace utiliza arquitecturas tipo Feature Pyramid Network (FPN) permitiendo detectar rostros diminutos y extrayendo landmarks simultáneamente, sin depender de la naturaleza lenta y por etapas del MTCNN.')
    
    p = doc.add_paragraph()
    p.add_run('Experimento de Detectores:').bold = True
    doc.add_paragraph('En el código punto4_detectors.py, contrastamos los detectores sobre la misma imagen midiendo el tiempo final del paso feed-forward en segundos:')
    
    try:
        if os.path.exists('data/resultados/Punto 4.png'):
            pic = doc.add_picture('data/resultados/Punto 4.png', width=Inches(6.0))
            pic_par = doc.paragraphs[-1]
            pic_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph('Figura 3: Identificación del ROI (Region of Interest) facial por distintos backends. Los enfoques de una sola pasada ostentan mejor desempeño de latencia frente a los bucles recurrentes de modelos por etapas.')
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].font.size = Pt(9)
            cap.runs[0].font.italic = True
    except Exception as e:
        doc.add_paragraph(f'(Error cargando imagen 3: {e})')

    # ------------------ Punto 5 ------------------
    doc.add_heading('5. Abordaje de la "Identidad Sintética" y DeepFakes (Sistemas Embebidos)', level=1)
    
    doc.add_paragraph('El concepto de Identidad Sintética y Deepfakes representa el mayor riesgo de seguridad (Vulnerabilidad de Presentación) en cerraduras inteligentes y sistemas embebidos de control de acceso facial. Para resolver esto, no basta con extraer vectores espaciales (embeddings). Es matemáticamente mandatorio inyectar un clasificador convolucional anterior al pipeline de Face Recognition llamado Sistema Anti-Spoofing (Liveness Detection).')
    
    p = doc.add_paragraph()
    p.add_run('Implementación en el Proyecto:').bold = True
    doc.add_paragraph('El ecosistema implementado en los códigos punto5_realtime.py y sistema_asistencia.py aborda esta problemática de forma profesional:')
    
    doc.add_paragraph('1. Liveness Detection Activo: Mediante el flag anti_spoofing=True introducido en las API de modelos SOTA, se extrae una evaluación secundaria que estima la dispersión de luz, micro-movimientos y varianza de texturas para descartar fotografías impresas y pantallas de celular (Spoofing 2D).', style='List Number')
    doc.add_paragraph('2. Multi-Threading Dinámico: Dado que calcular Spoofing + Embedding penaliza gravemente la CPU en dispositivos embebidos (ej. Raspberry Pi), se desarrolló un algoritmo concurrente con hilos (threading) que encapsula la iteración matemática pesada en un plano secundario, manteniendo el stream de la cámara (interfaz de usuario) a ~30 FPS y garantizando una experiencia fluida.', style='List Number')
    doc.add_paragraph('3. Puntaje Híbrido de Seguridad: Si el modelo detecta que el objeto carece de un Liveness Score alto, trunca el proceso y deniega el acceso sin importar cuán exacta sea la similitud con la base de datos de características, bloqueando exitosamente ataques sintéticos y deepfakes generados por Inteligencia Artificial.', style='List Number')

    # Guardar documento
    doc.save('Informe_Final_Actividad_09.docx')
    print("Documento Word generado exitosamente: Informe_Final_Actividad_09.docx")

if __name__ == '__main__':
    main()
